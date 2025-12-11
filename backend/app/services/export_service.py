"""
Export Service - Generate PDF, DOCX, CSV exports of contract data
"""

from typing import List, Dict, Any, Optional
from datetime import datetime
from io import BytesIO
import csv

from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, PageBreak
from reportlab.platypus import Image as RLImage
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class ExportService:
    """Handle export operations for contracts, reports, and data"""

    @staticmethod
    def generate_contract_pdf(contract_data: Dict[str, Any]) -> BytesIO:
        """Generate PDF report for contract analysis"""
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()

        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a2e1a'),
            spaceAfter=30,
            alignment=TA_CENTER
        )

        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#3d2f28'),
            spaceAfter=12,
            spaceBefore=12
        )

        # Title
        story.append(Paragraph("Contract Analysis Report", title_style))
        story.append(Spacer(1, 0.2 * inch))

        # Metadata
        metadata = [
            ["Contract Name:", contract_data.get('contract_name', 'N/A')],
            ["Analysis Date:", datetime.now().strftime("%B %d, %Y at %I:%M %p")],
            ["Risk Level:", contract_data.get('risk_level', 'Unknown')],
            ["Risk Score:", f"{contract_data.get('risk_score', 0)}/100"],
        ]

        metadata_table = Table(metadata, colWidths=[2 * inch, 4 * inch])
        metadata_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f5edd8')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#1a2e1a')),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        story.append(metadata_table)
        story.append(Spacer(1, 0.3 * inch))

        # Summary
        story.append(Paragraph("Executive Summary", heading_style))
        summary_text = contract_data.get('summary', 'No summary available.')
        story.append(Paragraph(summary_text, styles['BodyText']))
        story.append(Spacer(1, 0.2 * inch))

        # Key Terms
        if contract_data.get('key_terms'):
            story.append(Paragraph("Key Terms", heading_style))
            for term in contract_data['key_terms']:
                story.append(Paragraph(f"• {term}", styles['BodyText']))
            story.append(Spacer(1, 0.2 * inch))

        # Clauses
        if contract_data.get('clauses'):
            story.append(Paragraph("Detected Clauses", heading_style))
            clause_data = [["Type", "Content", "Risk"]]
            for clause in contract_data['clauses'][:10]:  # Limit to 10 for PDF
                clause_data.append([
                    clause.get('type', 'Unknown'),
                    clause.get('text', '')[:100] + '...',  # Truncate long text
                    clause.get('risk_level', 'Low')
                ])

            clause_table = Table(clause_data, colWidths=[1.5 * inch, 3.5 * inch, 1 * inch])
            clause_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3d2f28')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
            ]))
            story.append(clause_table)
            story.append(Spacer(1, 0.2 * inch))

        # Risks
        if contract_data.get('risks'):
            story.append(Paragraph("Identified Risks", heading_style))
            for risk in contract_data['risks']:
                severity = risk.get('severity', 'Unknown')
                color = {
                    'Critical': 'red',
                    'High': 'orange',
                    'Medium': 'yellow',
                    'Low': 'green'
                }.get(severity, 'black')

                risk_text = f"<font color='{color}'><b>[{severity}]</b></font> {risk.get('description', '')}"
                story.append(Paragraph(risk_text, styles['BodyText']))
            story.append(Spacer(1, 0.2 * inch))

        # Footer
        story.append(Spacer(1, 0.5 * inch))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        story.append(Paragraph(
            "Generated by DafLegal - AI-Powered Legal Analysis Platform<br/>www.daflegal.com",
            footer_style
        ))

        doc.build(story)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_compliance_pdf(compliance_data: Dict[str, Any]) -> BytesIO:
        """Generate PDF report for compliance check"""
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()

        # Title
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1a2e1a'),
            spaceAfter=30,
            alignment=TA_CENTER
        )
        story.append(Paragraph("Compliance Check Report", title_style))
        story.append(Spacer(1, 0.2 * inch))

        # Metadata
        metadata = [
            ["Playbook:", compliance_data.get('playbook_name', 'N/A')],
            ["Check Date:", datetime.now().strftime("%B %d, %Y at %I:%M %p")],
            ["Compliance Score:", f"{compliance_data.get('score', 0)}%"],
            ["Status:", compliance_data.get('status', 'Unknown')],
        ]

        metadata_table = Table(metadata, colWidths=[2 * inch, 4 * inch])
        metadata_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f5edd8')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#1a2e1a')),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        story.append(metadata_table)
        story.append(Spacer(1, 0.3 * inch))

        # Issues
        if compliance_data.get('issues'):
            heading_style = ParagraphStyle(
                'Heading',
                parent=styles['Heading2'],
                fontSize=16,
                textColor=colors.HexColor('#3d2f28'),
                spaceAfter=12
            )
            story.append(Paragraph("Compliance Issues", heading_style))

            issue_data = [["Severity", "Rule", "Description"]]
            for issue in compliance_data['issues']:
                issue_data.append([
                    issue.get('severity', 'Unknown'),
                    issue.get('rule_name', ''),
                    issue.get('description', '')[:150] + '...'
                ])

            issue_table = Table(issue_data, colWidths=[1 * inch, 2 * inch, 3 * inch])
            issue_table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3d2f28')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, 0), 11),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('FONTSIZE', (0, 1), (-1, -1), 9),
            ]))
            story.append(issue_table)

        doc.build(story)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_clauses_docx(clauses: List[Dict[str, Any]]) -> BytesIO:
        """Generate DOCX document with clauses"""
        document = Document()
        buffer = BytesIO()

        # Title
        title = document.add_heading('Clause Library Export', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Metadata
        document.add_paragraph(f"Exported: {datetime.now().strftime('%B %d, %Y at %I:%M %p')}")
        document.add_paragraph(f"Total Clauses: {len(clauses)}")
        document.add_paragraph()  # Blank line

        # Clauses
        for i, clause in enumerate(clauses, 1):
            # Clause heading
            heading = document.add_heading(f"Clause {i}: {clause.get('name', 'Unnamed')}", level=2)

            # Type
            type_para = document.add_paragraph()
            type_para.add_run('Type: ').bold = True
            type_para.add_run(clause.get('type', 'Unknown'))

            # Category
            if clause.get('category'):
                cat_para = document.add_paragraph()
                cat_para.add_run('Category: ').bold = True
                cat_para.add_run(clause['category'])

            # Content
            content_para = document.add_paragraph()
            content_para.add_run('Content:').bold = True
            document.add_paragraph(clause.get('text', 'No content available'))

            # Tags
            if clause.get('tags'):
                tags_para = document.add_paragraph()
                tags_para.add_run('Tags: ').bold = True
                tags_para.add_run(', '.join(clause['tags']))

            # Separator
            document.add_paragraph('_' * 80)

        document.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_analytics_csv(analytics_data: List[Dict[str, Any]]) -> BytesIO:
        """Generate CSV file with analytics data"""
        buffer = BytesIO()

        if not analytics_data:
            return buffer

        # Get headers from first row
        headers = list(analytics_data[0].keys())

        # Write to CSV
        output = buffer
        output.write(','.join(headers).encode('utf-8'))
        output.write(b'\n')

        for row in analytics_data:
            values = [str(row.get(h, '')) for h in headers]
            output.write(','.join(values).encode('utf-8'))
            output.write(b'\n')

        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_contracts_csv(contracts: List[Dict[str, Any]]) -> BytesIO:
        """Generate CSV export of contracts list"""
        buffer = BytesIO()

        headers = ['ID', 'Name', 'Upload Date', 'Risk Level', 'Risk Score', 'Status']

        output = buffer
        output.write(','.join(headers).encode('utf-8'))
        output.write(b'\n')

        for contract in contracts:
            row = [
                str(contract.get('id', '')),
                contract.get('contract_name', ''),
                contract.get('created_at', ''),
                contract.get('risk_level', ''),
                str(contract.get('risk_score', '')),
                contract.get('status', '')
            ]
            # Escape commas and quotes in values
            row = [f'"{v.replace(""", """""")}"' if ',' in v or '"' in v else v for v in row]
            output.write(','.join(row).encode('utf-8'))
            output.write(b'\n')

        buffer.seek(0)
        return buffer
